from __future__ import annotations

import csv
import json
import logging
import math
import re
from dataclasses import dataclass
from datetime import datetime, timezone
from io import TextIOWrapper
from pathlib import Path
from tempfile import TemporaryDirectory
from threading import Lock
from typing import Any, Callable
from uuid import uuid4
from zipfile import BadZipFile, ZipFile

from matplotlib.backends.backend_agg import FigureCanvasAgg
from matplotlib.figure import Figure
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


_DATA_ARCHIVE = (
	Path(__file__).resolve().parents[3]
	/ "data"
	/ "eu_resilience"
	/ "documents"
	/ "data.zip"
)
_PLAYBOOK_NAME = "EU Cross-Agency Coordination Playbook"
_PLAYBOOK_VERSION = "1.0-demo"
_REQUIRED_CRITERIA = 3
_CASE_REGISTER: dict[str, dict[str, Any]] = {}
_CASE_REGISTER_LOCK = Lock()
_REPORT_TEMP_DIRECTORY = TemporaryDirectory(prefix="cosmopilot-eu-reports-")
_REPORT_OUTPUT_DIR = Path(_REPORT_TEMP_DIRECTORY.name)
_REPORT_CONTENT_TYPE = (
	"application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
_MAX_GENERATED_DOCUMENTS = 64
_DOCUMENT_ID_PATTERN = re.compile(r"^[0-9a-f]{32}$")
_GENERATED_DOCUMENTS: dict[str, "GeneratedDocument"] = {}
_GENERATED_DOCUMENTS_LOCK = Lock()
_DOMAIN_LABELS = (
	("environmental_risk", "Environmental"),
	("population_vulnerability", "Population"),
	("health_pressure", "Health"),
	("food_risk", "Food"),
	("economic_stress", "Economic"),
	("financial_exposure", "Financial"),
)
logger = logging.getLogger("eu_resilience_functions")


@dataclass(frozen=True, slots=True)
class GeneratedDocument:
	path: Path
	file_name: str
	title: str
	content_type: str


def get_generated_document(document_id: str) -> GeneratedDocument | None:
	"""Return a generated report only for an exact opaque document identifier."""
	if not isinstance(document_id, str) or not _DOCUMENT_ID_PATTERN.fullmatch(
		document_id
	):
		return None

	with _GENERATED_DOCUMENTS_LOCK:
		document = _GENERATED_DOCUMENTS.get(document_id)
		if document is not None and not document.path.is_file():
			_GENERATED_DOCUMENTS.pop(document_id, None)
			return None
		return document


def _register_generated_document(
	document_id: str,
	document: GeneratedDocument,
) -> None:
	with _GENERATED_DOCUMENTS_LOCK:
		while len(_GENERATED_DOCUMENTS) >= _MAX_GENERATED_DOCUMENTS:
			oldest_id = next(iter(_GENERATED_DOCUMENTS))
			oldest = _GENERATED_DOCUMENTS.pop(oldest_id)
			try:
				oldest.path.unlink()
			except OSError as exc:
				logger.warning(
					"Unable to remove expired generated document %s: %s",
					oldest.path,
					exc,
				)
		_GENERATED_DOCUMENTS[document_id] = document


def _coerce_csv_value(value: str | None) -> Any:
	if value is None:
		return None

	value = value.strip()
	if not value:
		return None

	try:
		return int(value)
	except ValueError:
		try:
			return float(value)
		except ValueError:
			return value


def _read_csv(member: str) -> list[dict[str, Any]]:
	try:
		with ZipFile(_DATA_ARCHIVE) as archive:
			with archive.open(member) as raw_stream:
				with TextIOWrapper(raw_stream, encoding="utf-8-sig", newline="") as stream:
					return [
						{
							key: _coerce_csv_value(value)
							for key, value in row.items()
							if key is not None
						}
						for row in csv.DictReader(stream)
					]
	except (FileNotFoundError, KeyError, BadZipFile) as exc:
		raise RuntimeError(
			f"Unable to read resilience evidence member {member!r} from {_DATA_ARCHIVE}."
		) from exc


def _resolve_country(country: str) -> dict[str, Any]:
	if not isinstance(country, str) or not country.strip():
		raise ValueError("country must be a non-empty EU country name or ISO code")

	lookup = country.strip().casefold()
	aliases = {
		"czechia": "CZE",
		"gr": "GRC",
	}
	lookup = aliases.get(lookup, lookup).casefold()

	for region in _read_csv("data/regions.csv"):
		identifiers = {
			str(region[field]).casefold()
			for field in ("country", "iso2", "iso3")
			if region.get(field)
		}
		if lookup in identifiers:
			return region

	raise ValueError(f"Unknown EU27 country or ISO code: {country!r}")


def _country_rows(member: str, iso3: str) -> list[dict[str, Any]]:
	return [row for row in _read_csv(member) if row.get("iso3") == iso3]


def _single_country_row(member: str, iso3: str) -> dict[str, Any] | None:
	rows = _country_rows(member, iso3)
	if len(rows) > 1:
		raise RuntimeError(f"Expected one {member} record for {iso3}, found {len(rows)}.")
	return rows[0] if rows else None


def _evidence_package_id(scorecard: dict[str, Any]) -> str:
	date = str(scorecard["date"]).replace("-", "")
	return f"EU27-{date}-{scorecard['iso3']}"


def _evidence_section(
	record: dict[str, Any],
	*,
	source_label: str,
	classification: str,
	snapshot_date: str,
) -> dict[str, Any]:
	identity_fields = {"date", "iso2", "iso3", "country", "zone"}
	return {
		"source_label": source_label,
		"classification": classification,
		"observation_date": record.get("date") or snapshot_date,
		"geographic_grain": "country",
		"values": {
			key: value
			for key, value in record.items()
			if key not in identity_fields
		},
	}


def get_resilience_priorities(limit: int) -> dict[str, Any]:
	"""Return the highest-risk countries in the static EU27 demo snapshot."""
	if isinstance(limit, bool) or not isinstance(limit, int):
		raise TypeError("limit must be an integer")
	if not 1 <= limit <= 27:
		raise ValueError("limit must be between 1 and 27")

	scorecard = _read_csv("data/region_scorecard.csv")
	ranked = sorted(
		scorecard,
		key=lambda row: (-float(row["overall_risk"]), str(row["country"])),
	)
	snapshot_dates = sorted({str(row["date"]) for row in ranked})
	limitations = [
		"The values are a static demo snapshot, not live operational data.",
		"Overall and domain risk scores are synthetic derived composites.",
		"All observations have country-level geographic grain.",
	]
	if len(snapshot_dates) != 1:
		limitations.append(
			"The scorecard contains conflicting observation dates: "
			+ ", ".join(snapshot_dates)
		)

	priorities = []
	for rank, row in enumerate(ranked[:limit], start=1):
		priorities.append(
			{
				"rank": rank,
				"country": row["country"],
				"iso3": row["iso3"],
				"zone": row["zone"],
				"priority_score": row["overall_risk"],
				"severity": row["severity"],
				"observation_date": row["date"],
				"evidence_package_id": _evidence_package_id(row),
				"lead_response_agency": row["lead_response_agency"],
				"classification": "synthetic-derived composite",
				"domain_scores": {
					"environmental_risk": row["environmental_risk"],
					"population_vulnerability": row["population_vulnerability"],
					"health_pressure": row["health_pressure"],
					"food_risk": row["food_risk"],
					"economic_stress": row["economic_stress"],
					"financial_exposure": row["financial_exposure"],
				},
			}
		)

	return {
		"snapshot_date": snapshot_dates[0] if len(snapshot_dates) == 1 else None,
		"snapshot_type": "static demo snapshot",
		"source_label": "Curated EU27 resilience scorecard",
		"geographic_grain": "country",
		"ranking_metric": "overall_risk descending",
		"available_country_count": len(ranked),
		"returned_country_count": len(priorities),
		"priorities": priorities,
		"limitations": limitations,
	}


def get_country_resilience_evidence(country: str) -> dict[str, Any]:
	"""Return a joined, provenance-labelled evidence package for one EU country."""
	region = _resolve_country(country)
	iso3 = str(region["iso3"])
	scorecard = _single_country_row("data/region_scorecard.csv", iso3)
	if scorecard is None:
		raise RuntimeError(f"No resilience scorecard record is available for {iso3}.")

	snapshot_date = str(scorecard["date"])
	evidence: dict[str, Any] = {}
	missing_sources: list[str] = []
	section_specs = (
		(
			"environmental",
			"data/copernicus_environment_daily.csv",
			"Copernicus curated environmental observations",
			"source-derived",
		),
		(
			"wildfire",
			"data/effis_wildfire_country_2026.csv",
			"EFFIS curated wildfire observations",
			"source-derived",
		),
		(
			"food_supply",
			"data/efsa_food_supply_daily.csv",
			"EFSA curated food-supply observations",
			"source-derived",
		),
		(
			"health",
			"data/ecdc_health_daily.csv",
			"ECDC curated health observations",
			"source-derived",
		),
		(
			"economic",
			"data/esm_macro_stability.csv",
			"ESM curated macro-stability observations",
			"source-derived",
		),
		(
			"banking_exposure",
			"data/eba_banking_exposure.csv",
			"Curated EBA-aligned banking exposure",
			"synthetic",
		),
		(
			"banking_stress_test",
			"data/eba_2025_stress_test_country_synthetic.csv",
			"EBA 2025 aggregates with a synthetic country overlay",
			"mixed derived and synthetic",
		),
	)
	for section_name, member, source_label, classification in section_specs:
		record = _single_country_row(member, iso3)
		if record is None:
			missing_sources.append(source_label)
			continue
		evidence[section_name] = _evidence_section(
			record,
			source_label=source_label,
			classification=classification,
			snapshot_date=snapshot_date,
		)

	evidence["population"] = _evidence_section(
		region,
		source_label="Curated EU27 demographic reference",
		classification="mixed source-derived and synthetic composite",
		snapshot_date=snapshot_date,
	)

	alerts = []
	for alert in _country_rows("data/realtime_alerts.csv", iso3):
		alerts.append(
			{
				"event_id": alert["event_id"],
				"event_time": alert["event_time"],
				"source_label": alert["source_agency"],
				"alert_type": alert["alert_type"],
				"severity": alert["severity"],
				"confidence": alert["confidence"],
				"summary": alert["summary"],
				"classification": "synthetic",
			}
		)

	limitations = [
		f"Evidence is a static demo snapshot observed on {snapshot_date}.",
		"Geographic grain is country level; no regional or local precision is implied.",
		"Overall risk, domain scores, and population vulnerability are synthetic derived composites.",
		"Banking records and climate-credit losses are synthetic overlays.",
		"EBA adverse-scenario values are stress-test scenarios, not forecasts of bank failure or evidence of current distress.",
	]
	stress_test = evidence.get("banking_stress_test", {}).get("values", {})
	if str(stress_test.get("stress_data_basis", "")).startswith("Synthetic EU27 fallback"):
		limitations.append(
			"Banking stress values use a synthetic EU27 fallback because no EBA reference bank was available for this country."
		)
	if missing_sources:
		limitations.append("Missing evidence sources: " + ", ".join(missing_sources))

	return {
		"evidence_package_id": _evidence_package_id(scorecard),
		"snapshot_date": snapshot_date,
		"snapshot_type": "static demo snapshot",
		"country": {
			"name": region["country"],
			"iso2": region["iso2"],
			"iso3": iso3,
			"zone": region["zone"],
		},
		"geographic_grain": "country",
		"priority": {
			"severity": scorecard["severity"],
			"overall_risk": scorecard["overall_risk"],
			"lead_response_agency": scorecard["lead_response_agency"],
			"classification": "synthetic-derived composite",
			"domain_scores": {
				"environmental_risk": scorecard["environmental_risk"],
				"population_vulnerability": scorecard["population_vulnerability"],
				"health_pressure": scorecard["health_pressure"],
				"food_risk": scorecard["food_risk"],
				"economic_stress": scorecard["economic_stress"],
				"financial_exposure": scorecard["financial_exposure"],
			},
		},
		"evidence": evidence,
		"alerts": alerts,
		"confidence": {
			"alert_confidence_values": [alert["confidence"] for alert in alerts],
			"note": "Alert confidence is source metadata, not model certainty or an independently calibrated package score.",
		},
		"classifications": {
			"source_derived": [
				"environmental observations",
				"wildfire observations",
				"food-supply observations",
				"health observations",
				"macroeconomic observations",
			],
			"derived": ["EBA country aggregates where reference banks are available"],
			"synthetic": [
				"alerts",
				"banking records",
				"climate-credit losses",
				"composite vulnerability and risk scores",
				"country fallback estimates where disclosed",
			],
		},
		"limitations": limitations,
	}


def _metric_label(metric: str) -> str:
	return metric.replace("_", " ").replace("pct", "%").replace("eur", "EUR").title()


def _format_metric_value(metric: str, value: Any) -> str:
	if not isinstance(value, (int, float)) or isinstance(value, bool):
		return str(value)
	if metric.endswith("_eur_m"):
		return f"EUR {value:,.1f}m"
	if metric.endswith("_eur_bn"):
		return f"EUR {value:,.1f}bn"
	if metric.endswith("_ha"):
		return f"{value:,.0f} ha"
	if metric.endswith("_bps"):
		return f"{value:,.0f} bps"
	if metric.endswith("_pct") or metric.endswith("_percentage"):
		return f"{value:,.1f}%"
	if isinstance(value, float) and not value.is_integer():
		return f"{value:,.2f}"
	return f"{value:,.0f}"


def _strongest_domains(domain_scores: dict[str, Any]) -> list[tuple[str, float]]:
	scores = [
		(label, float(domain_scores[metric]))
		for metric, label in _DOMAIN_LABELS
		if isinstance(domain_scores.get(metric), (int, float))
		and not isinstance(domain_scores.get(metric), bool)
	]
	return sorted(scores, key=lambda item: (-item[1], item[0]))


def _set_cell_shading(cell: Any, fill: str) -> None:
	properties = cell._tc.get_or_add_tcPr()
	shading = properties.find(qn("w:shd"))
	if shading is None:
		shading = OxmlElement("w:shd")
		properties.append(shading)
	shading.set(qn("w:fill"), fill)


def _style_table_header(table: Any) -> None:
	for cell in table.rows[0].cells:
		_set_cell_shading(cell, "174895")
		for run in cell.paragraphs[0].runs:
			run.font.bold = True
			run.font.color.rgb = RGBColor(255, 255, 255)


def _configure_report_document(document: Any, title: str) -> None:
	document.core_properties.title = title
	document.core_properties.subject = "EU27 resilience evidence briefing"
	document.core_properties.author = "EU Resilience Desk"
	document.core_properties.keywords = (
		"EU27, resilience, coordination, synthetic demo data"
	)

	for section in document.sections:
		section.top_margin = Inches(0.65)
		section.bottom_margin = Inches(0.65)
		section.left_margin = Inches(0.72)
		section.right_margin = Inches(0.72)
		header = section.header.paragraphs[0]
		header.text = "EU RESILIENCE DESK  /  CONTROLLED DEMO BRIEFING"
		header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
		for run in header.runs:
			run.font.name = "Aptos"
			run.font.size = Pt(8)
			run.font.bold = True
			run.font.color.rgb = RGBColor(100, 113, 108)
		footer = section.footer.paragraphs[0]
		footer.text = (
			"Country-level static demo snapshot. Synthetic and derived values "
			"are disclosed in this document."
		)
		footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
		for run in footer.runs:
			run.font.name = "Aptos"
			run.font.size = Pt(8)
			run.font.color.rgb = RGBColor(100, 113, 108)

	normal = document.styles["Normal"]
	normal.font.name = "Aptos"
	normal.font.size = Pt(10)
	normal.paragraph_format.space_after = Pt(6)
	for style_name in ("Title", "Heading 1", "Heading 2"):
		style = document.styles[style_name]
		style.font.name = "Aptos Display"
		style.font.color.rgb = RGBColor(16, 54, 111)


def _add_report_cover(
	document: Any,
	*,
	title: str,
	subtitle: str,
	snapshot_date: str,
	scope: str,
	generated_at: str,
) -> None:
	banner = document.add_table(rows=1, cols=1)
	banner.alignment = WD_TABLE_ALIGNMENT.CENTER
	cell = banner.cell(0, 0)
	_set_cell_shading(cell, "174895")
	paragraph = cell.paragraphs[0]
	paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
	paragraph.paragraph_format.space_before = Pt(18)
	paragraph.paragraph_format.space_after = Pt(18)
	run = paragraph.add_run(title)
	run.font.name = "Aptos Display"
	run.font.size = Pt(25)
	run.font.bold = True
	run.font.color.rgb = RGBColor(255, 255, 255)

	subtitle_paragraph = document.add_paragraph()
	subtitle_paragraph.paragraph_format.space_before = Pt(10)
	subtitle_run = subtitle_paragraph.add_run(subtitle)
	subtitle_run.font.size = Pt(13)
	subtitle_run.font.bold = True
	subtitle_run.font.color.rgb = RGBColor(35, 113, 83)

	meta = document.add_table(rows=2, cols=3)
	meta.style = "Light Shading Accent 1"
	meta.alignment = WD_TABLE_ALIGNMENT.CENTER
	for index, heading in enumerate(("Snapshot", "Scope", "Generated")):
		meta.cell(0, index).text = heading
		meta.cell(1, index).text = (
			snapshot_date if index == 0 else scope if index == 1 else generated_at
		)
	_style_table_header(meta)
	document.add_paragraph(
		"Decision-support material for internal review. This report does not "
		"authorize public warnings, funding, financial intervention, or "
		"emergency action."
	)


def _save_country_charts(
	package: dict[str, Any],
	output_directory: Path,
) -> list[tuple[Path, str]]:
	domain_scores = package["priority"]["domain_scores"]
	labels = [label for _, label in _DOMAIN_LABELS]
	values = [float(domain_scores[metric]) for metric, _ in _DOMAIN_LABELS]

	radar_path = output_directory / "country-risk-radar.png"
	angles = [
		index / float(len(labels)) * 2 * math.pi
		for index in range(len(labels))
	]
	closed_angles = angles + angles[:1]
	closed_values = values + values[:1]
	figure = Figure(figsize=(7.4, 5.2), facecolor="white")
	FigureCanvasAgg(figure)
	axis = figure.subplots(subplot_kw={"polar": True})
	axis.set_theta_offset(math.pi / 2)
	axis.set_theta_direction(-1)
	axis.set_xticks(angles)
	axis.set_xticklabels(labels, fontsize=9, color="#16231f")
	axis.set_ylim(0, 100)
	axis.set_yticks((20, 40, 60, 80, 100))
	axis.set_yticklabels(("20", "40", "60", "80", "100"), fontsize=7)
	axis.plot(closed_angles, closed_values, color="#174895", linewidth=2.5)
	axis.fill(closed_angles, closed_values, color="#174895", alpha=0.2)
	axis.scatter(angles, values, color="#c94f3d", s=34, zorder=3)
	axis.set_title(
		f"{package['country']['name']} domain risk profile",
		pad=22,
		fontweight="bold",
		color="#10366f",
	)
	figure.tight_layout()
	figure.savefig(radar_path, dpi=190, bbox_inches="tight", facecolor="white")

	bar_path = output_directory / "country-risk-bars.png"
	ranked = _strongest_domains(domain_scores)
	bar_labels = [label for label, _ in reversed(ranked)]
	bar_values = [score for _, score in reversed(ranked)]
	figure = Figure(figsize=(7.4, 4.2), facecolor="white")
	FigureCanvasAgg(figure)
	axis = figure.subplots()
	colors = [
		"#c94f3d" if score >= 75 else "#b77b17" if score >= 60 else "#237153"
		for score in bar_values
	]
	bars = axis.barh(bar_labels, bar_values, color=colors, height=0.62)
	axis.set_xlim(0, 100)
	axis.set_xlabel("Derived risk score (0-100)")
	axis.set_title(
		"Decision drivers ranked by pressure",
		fontweight="bold",
		color="#10366f",
	)
	axis.grid(axis="x", alpha=0.18)
	axis.spines[["top", "right", "left"]].set_visible(False)
	for bar, value in zip(bars, bar_values):
		axis.text(
			min(value + 1.5, 96),
			bar.get_y() + bar.get_height() / 2,
			f"{value:.0f}",
			va="center",
			fontweight="bold",
			fontsize=9,
		)
	figure.tight_layout()
	figure.savefig(bar_path, dpi=190, bbox_inches="tight", facecolor="white")

	return [
		(radar_path, "Figure 1. Six-domain country risk profile."),
		(bar_path, "Figure 2. Ranked derived decision drivers."),
	]


def _save_priority_charts(
	priorities: dict[str, Any],
	output_directory: Path,
) -> list[tuple[Path, str]]:
	rows = priorities["priorities"]
	bar_path = output_directory / "eu-priority-ranking.png"
	figure = Figure(
		figsize=(7.4, max(3.4, len(rows) * 0.52 + 1.5)),
		facecolor="white",
	)
	FigureCanvasAgg(figure)
	axis = figure.subplots()
	countries = [str(row["country"]) for row in reversed(rows)]
	scores = [float(row["priority_score"]) for row in reversed(rows)]
	bars = axis.barh(countries, scores, color="#174895", height=0.6)
	axis.set_xlim(0, 100)
	axis.set_xlabel("Overall derived priority score (0-100)")
	axis.set_title(
		"EU27 leadership priority ranking",
		fontweight="bold",
		color="#10366f",
	)
	axis.grid(axis="x", alpha=0.18)
	axis.spines[["top", "right", "left"]].set_visible(False)
	for bar, value in zip(bars, scores):
		axis.text(
			min(value + 1.4, 96),
			bar.get_y() + bar.get_height() / 2,
			f"{value:.0f}",
			va="center",
			fontweight="bold",
			fontsize=9,
		)
	figure.tight_layout()
	figure.savefig(bar_path, dpi=190, bbox_inches="tight", facecolor="white")

	heatmap_path = output_directory / "eu-domain-heatmap.png"
	matrix = [
		[float(row["domain_scores"][metric]) for metric, _ in _DOMAIN_LABELS]
		for row in rows
	]
	figure = Figure(
		figsize=(7.4, max(3.5, len(rows) * 0.5 + 1.8)),
		facecolor="white",
	)
	FigureCanvasAgg(figure)
	axis = figure.subplots()
	image = axis.imshow(matrix, cmap="YlOrRd", vmin=0, vmax=100, aspect="auto")
	axis.set_xticks(range(len(_DOMAIN_LABELS)))
	axis.set_xticklabels(
		[label for _, label in _DOMAIN_LABELS],
		rotation=30,
		ha="right",
		fontsize=8,
	)
	axis.set_yticks(range(len(rows)))
	axis.set_yticklabels([str(row["country"]) for row in rows], fontsize=9)
	axis.set_title(
		"Cross-domain pressure heatmap",
		fontweight="bold",
		color="#10366f",
	)
	for row_index, row_values in enumerate(matrix):
		for column_index, value in enumerate(row_values):
			axis.text(
				column_index,
				row_index,
				f"{value:.0f}",
				ha="center",
				va="center",
				fontsize=8,
				fontweight="bold",
				color="white" if value >= 62 else "#16231f",
			)
	colorbar = figure.colorbar(image, ax=axis, fraction=0.028, pad=0.03)
	colorbar.set_label("Derived score", fontsize=8)
	figure.tight_layout()
	figure.savefig(
		heatmap_path,
		dpi=190,
		bbox_inches="tight",
		facecolor="white",
	)

	return [
		(bar_path, "Figure 1. Ranked EU27 leadership priorities."),
		(heatmap_path, "Figure 2. Comparative domain-pressure heatmap."),
	]


def _add_chart_pack(
	document: Any,
	charts: list[tuple[Path, str]],
) -> None:
	document.add_heading("Visual evidence", level=1)
	for chart_path, caption in charts:
		document.add_picture(str(chart_path), width=Inches(6.35))
		paragraph = document.add_paragraph()
		paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
		run = paragraph.add_run(caption)
		run.font.italic = True
		run.font.size = Pt(8)
		run.font.color.rgb = RGBColor(100, 113, 108)


def _add_limitations(document: Any, limitations: list[str]) -> None:
	document.add_heading("Evidence boundaries", level=1)
	for limitation in limitations:
		document.add_paragraph(str(limitation), style="List Bullet")


def _build_country_report(
	package: dict[str, Any],
	charts: list[tuple[Path, str]],
	generated_at: str,
) -> Any:
	country = str(package["country"]["name"])
	title = f"{country} resilience decision briefing"
	document = Document()
	_configure_report_document(document, title)
	_add_report_cover(
		document,
		title=title,
		subtitle=(
			f"{package['priority']['severity']} priority | "
			f"Overall risk {package['priority']['overall_risk']}/100"
		),
		snapshot_date=str(package["snapshot_date"]),
		scope=f"{country} | country-level",
		generated_at=generated_at,
	)

	document.add_heading("Executive assessment", level=1)
	document.add_paragraph(
		f"{country} is classified as {package['priority']['severity']} priority "
		f"with an overall derived risk score of "
		f"{package['priority']['overall_risk']}/100. "
		f"{package['priority']['lead_response_agency']} is the indicated lead "
		"response agency for internal review."
	)
	document.add_heading("Principal decision drivers", level=2)
	for label, score in _strongest_domains(
		package["priority"]["domain_scores"]
	)[:4]:
		document.add_paragraph(
			f"{label}: {score:.0f}/100",
			style="List Bullet",
		)

	_add_chart_pack(document, charts)

	document.add_heading("Evidence register", level=1)
	table = document.add_table(rows=1, cols=4)
	table.style = "Light Shading Accent 1"
	table.alignment = WD_TABLE_ALIGNMENT.CENTER
	for cell, heading in zip(
		table.rows[0].cells,
		("Evidence area", "Classification", "Observed", "Selected indicators"),
	):
		cell.text = heading
	_style_table_header(table)
	for section_name, section in package["evidence"].items():
		row = table.add_row().cells
		row[0].text = section_name.replace("_", " ").title()
		row[1].text = str(section["classification"])
		row[2].text = str(section["observation_date"])
		values = section.get("values", {})
		priority_metrics = [
			metric
			for metric in values
			if metric.endswith(
				("_risk", "_pressure", "_stress", "_exposure", "_pct")
			)
		]
		selected_metrics = (priority_metrics + list(values))[:3]
		row[3].text = "; ".join(
			f"{_metric_label(metric)}: "
			f"{_format_metric_value(metric, values[metric])}"
			for metric in dict.fromkeys(selected_metrics)
		)

	if package["alerts"]:
		document.add_heading("Active snapshot alerts", level=1)
		for alert in package["alerts"]:
			document.add_paragraph(
				f"{alert['alert_type']} | {alert['severity']} | "
				f"{alert['source_label']} | confidence {alert['confidence']}: "
				f"{alert['summary']}",
				style="List Bullet",
			)

	_add_limitations(document, package["limitations"])
	return document


def _build_priority_report(
	priorities: dict[str, Any],
	charts: list[tuple[Path, str]],
	generated_at: str,
) -> Any:
	title = "EU27 resilience leadership priorities"
	document = Document()
	_configure_report_document(document, title)
	_add_report_cover(
		document,
		title=title,
		subtitle=(
			f"Top {priorities['returned_country_count']} priorities from the "
			"curated resilience scorecard"
		),
		snapshot_date=str(priorities["snapshot_date"]),
		scope="EU27 comparative | country-level",
		generated_at=generated_at,
	)

	top = priorities["priorities"][0]
	document.add_heading("Executive assessment", level=1)
	document.add_paragraph(
		f"{top['country']} is the highest-ranked country in this snapshot with "
		f"an overall derived priority score of {top['priority_score']}/100. "
		"The ranking supports leadership triage only and does not constitute "
		"an emergency determination."
	)

	document.add_heading("Priority register", level=1)
	table = document.add_table(rows=1, cols=6)
	table.style = "Light Shading Accent 1"
	table.alignment = WD_TABLE_ALIGNMENT.CENTER
	for cell, heading in zip(
		table.rows[0].cells,
		("Rank", "Country", "Score", "Severity", "Strongest driver", "Lead"),
	):
		cell.text = heading
	_style_table_header(table)
	for priority in priorities["priorities"]:
		strongest_label, strongest_score = _strongest_domains(
			priority["domain_scores"]
		)[0]
		row = table.add_row().cells
		row[0].text = str(priority["rank"])
		row[1].text = str(priority["country"])
		row[2].text = f"{priority['priority_score']}/100"
		row[3].text = str(priority["severity"])
		row[4].text = f"{strongest_label} ({strongest_score:.0f})"
		row[5].text = str(priority["lead_response_agency"])

	_add_chart_pack(document, charts)
	_add_limitations(document, priorities["limitations"])
	return document


def generate_resilience_report(
	report_type: str,
	country: str | None,
	limit: int | None,
) -> dict[str, Any]:
	"""Generate a styled DOCX report from the authoritative local evidence."""
	if report_type not in {"country", "eu_priorities"}:
		raise ValueError("report_type must be 'country' or 'eu_priorities'")

	generated = datetime.now(timezone.utc)
	generated_at = generated.isoformat().replace("+00:00", "Z")
	_REPORT_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
	document_id = uuid4().hex

	with TemporaryDirectory(
		prefix=f"{document_id}-",
		dir=_REPORT_OUTPUT_DIR,
	) as chart_directory_name:
		chart_directory = Path(chart_directory_name)
		if report_type == "country":
			if not isinstance(country, str) or not country.strip():
				raise ValueError("country is required for a country report")
			if limit is not None:
				raise ValueError("limit must be null for a country report")
			data = get_country_resilience_evidence(country)
			canonical_country = str(data["country"]["name"])
			title = f"{canonical_country} resilience decision briefing"
			file_stem = re.sub(
				r"[^a-z0-9]+",
				"-",
				canonical_country.casefold(),
			).strip("-")
			charts = _save_country_charts(data, chart_directory)
			document = _build_country_report(data, charts, generated_at)
			scope = "country"
			snapshot_date = str(data["snapshot_date"])
			limitations = list(data["limitations"])
		else:
			if country is not None:
				raise ValueError("country must be null for an EU priorities report")
			if limit is None:
				limit = 5
			if isinstance(limit, bool) or not isinstance(limit, int):
				raise TypeError("limit must be an integer or null")
			if not 1 <= limit <= 10:
				raise ValueError("limit must be between 1 and 10")
			data = get_resilience_priorities(limit)
			title = "EU27 resilience leadership priorities"
			file_stem = f"eu27-top-{limit}-priorities"
			charts = _save_priority_charts(data, chart_directory)
			document = _build_priority_report(data, charts, generated_at)
			scope = "eu_priorities"
			snapshot_date = str(data["snapshot_date"])
			limitations = list(data["limitations"])

		file_name = (
			f"{file_stem}-{generated:%Y%m%d}-{document_id[:8]}.docx"
		)
		final_path = _REPORT_OUTPUT_DIR / file_name
		temporary_path = _REPORT_OUTPUT_DIR / f".{document_id}.tmp"
		try:
			document.save(temporary_path)
			temporary_path.replace(final_path)
		except OSError as exc:
			try:
				temporary_path.unlink()
			except FileNotFoundError:
				pass
			raise RuntimeError("Unable to write the generated DOCX report.") from exc

	generated_document = GeneratedDocument(
		path=final_path,
		file_name=file_name,
		title=title,
		content_type=_REPORT_CONTENT_TYPE,
	)
	_register_generated_document(document_id, generated_document)
	return {
		"status": "generated",
		"document": {
			"id": document_id,
			"title": title,
			"file_name": file_name,
			"download_url": f"/api/documents/{document_id}",
			"content_type": _REPORT_CONTENT_TYPE,
			"scope": scope,
			"generated_at": generated_at,
			"snapshot_date": snapshot_date,
			"chart_count": len(charts),
		},
		"limitations": limitations,
	}


def evaluate_coordination_playbook(
	country: str,
	evidence_package_id: str,
) -> dict[str, Any]:
	"""Apply the deterministic illustrative coordination criteria."""
	package = get_country_resilience_evidence(country)
	if evidence_package_id != package["evidence_package_id"]:
		raise ValueError(
			"evidence_package_id does not match the current country evidence package"
		)

	criterion_specs = (
		(
			"environmental_pressure",
			"Environmental risk",
			"environmental",
			"environmental_risk",
			70,
			"Copernicus",
		),
		(
			"health_system_pressure",
			"Health-system pressure",
			"health",
			"health_pressure",
			75,
			"ECDC",
		),
		(
			"food_system_pressure",
			"Food-system pressure",
			"food_supply",
			"food_risk",
			75,
			"EFSA",
		),
		(
			"financial_exposure",
			"Financial exposure",
			"banking_exposure",
			"financial_exposure",
			60,
			"EBA",
		),
		(
			"economic_stress",
			"Economic stress",
			"economic",
			"economic_stress",
			60,
			"ESM",
		),
	)

	criteria = []
	for criterion_id, name, section_name, metric, threshold, agency in criterion_specs:
		section = package["evidence"].get(section_name)
		value = section["values"].get(metric) if section else None
		if isinstance(value, (int, float)) and not isinstance(value, bool):
			status = "met" if value >= threshold else "not_met"
		else:
			status = "indeterminate"
		criteria.append(
			{
				"criterion_id": criterion_id,
				"name": name,
				"status": status,
				"observed_value": value,
				"threshold": threshold,
				"comparison": "greater_than_or_equal",
				"responsible_agency": agency,
				"source_label": section["source_label"] if section else None,
				"observation_date": section["observation_date"] if section else None,
			}
		)

	criteria_met = [item for item in criteria if item["status"] == "met"]
	criteria_not_met = [item for item in criteria if item["status"] == "not_met"]
	criteria_indeterminate = [item for item in criteria if item["status"] == "indeterminate"]
	if len(criteria_met) >= _REQUIRED_CRITERIA:
		eligibility = "eligible"
		coordination_eligible: bool | None = True
	elif len(criteria_met) + len(criteria_indeterminate) >= _REQUIRED_CRITERIA:
		eligibility = "indeterminate"
		coordination_eligible = None
	else:
		eligibility = "not_eligible"
		coordination_eligible = False

	lead_agency = (
		str(package["priority"]["lead_response_agency"])
		if coordination_eligible
		else None
	)
	participating_agencies = []
	if lead_agency:
		for criterion in criteria_met:
			agency = str(criterion["responsible_agency"])
			if agency != lead_agency and agency not in participating_agencies:
				participating_agencies.append(agency)

	review_steps = []
	if lead_agency:
		review_steps = [
			"Validate the evidence package, provenance, and disclosed limitations.",
			f"Convene {lead_agency} and the participating agencies for an internal coordination review.",
			"Confirm the country-level observations with the responsible national authorities.",
			"Record the coordination decision, follow-up owners, and review date.",
		]

	limitations = list(package["limitations"])
	limitations.append(
		"This is illustrative demo logic, not formal EU policy, legal advice, or an emergency-response mandate."
	)

	return {
		"playbook_name": _PLAYBOOK_NAME,
		"playbook_version": _PLAYBOOK_VERSION,
		"evidence_package_id": package["evidence_package_id"],
		"snapshot_date": package["snapshot_date"],
		"country": package["country"],
		"priority_level": package["priority"]["severity"],
		"eligibility": eligibility,
		"coordination_eligible": coordination_eligible,
		"required_criteria_count": _REQUIRED_CRITERIA,
		"criteria": criteria,
		"criteria_met": [item["criterion_id"] for item in criteria_met],
		"criteria_not_met": [item["criterion_id"] for item in criteria_not_met],
		"criteria_indeterminate": [
			item["criterion_id"] for item in criteria_indeterminate
		],
		"lead_agency": lead_agency,
		"participating_agencies": participating_agencies,
		"recommended_review_steps": review_steps,
		"limitations": limitations,
	}


def open_coordination_case(
	country: str,
	evidence_package_id: str,
	playbook_version: str,
	reason: str,
	lead_agency: str,
	participating_agencies: list[str],
	review_steps: list[str],
) -> dict[str, Any]:
	"""Record a validated, application-approved decision-card payload in memory."""
	if not isinstance(reason, str) or not reason.strip():
		raise ValueError("reason must be a non-empty string")
	if not isinstance(lead_agency, str) or not lead_agency.strip():
		raise ValueError("lead_agency must be a non-empty string")
	if not isinstance(participating_agencies, list) or not all(
		isinstance(agency, str) and agency.strip()
		for agency in participating_agencies
	):
		raise ValueError("participating_agencies must be a list of non-empty strings")
	if not isinstance(review_steps, list) or not all(
		isinstance(step, str) and step.strip()
		for step in review_steps
	):
		raise ValueError("review_steps must be a list of non-empty strings")

	evaluation = evaluate_coordination_playbook(country, evidence_package_id)
	if evaluation["coordination_eligible"] is not True:
		raise ValueError("the current evidence package is not eligible for coordination")
	if playbook_version != evaluation["playbook_version"]:
		raise ValueError("playbook_version does not match the current playbook evaluation")
	if lead_agency.casefold() != str(evaluation["lead_agency"]).casefold():
		raise ValueError("lead_agency does not match the current playbook evaluation")

	expected_agencies = [str(agency) for agency in evaluation["participating_agencies"]]
	provided_agencies = [agency.strip() for agency in participating_agencies]
	if len({agency.casefold() for agency in provided_agencies}) != len(provided_agencies):
		raise ValueError("participating_agencies must not contain duplicates")
	if {agency.casefold() for agency in provided_agencies} != {
		agency.casefold() for agency in expected_agencies
	}:
		raise ValueError(
			"participating_agencies do not match the current playbook evaluation"
		)

	expected_steps = [str(step) for step in evaluation["recommended_review_steps"]]
	if [" ".join(step.split()).casefold() for step in review_steps] != [
		" ".join(step.split()).casefold() for step in expected_steps
	]:
		raise ValueError("review_steps do not match the current playbook evaluation")

	submitted_at = datetime.now(timezone.utc).isoformat().replace("+00:00", "Z")
	case_id = f"EU-CRC-{datetime.now(timezone.utc):%Y%m%d}-{uuid4().hex[:8].upper()}"
	canonical_country = str(evaluation["country"]["name"])
	receipt = {
		"coordination_case_id": case_id,
		"submission_status": "submitted",
		"submitted_at": submitted_at,
		"country": canonical_country,
		"lead_agency": evaluation["lead_agency"],
		"participating_agencies": expected_agencies,
		"evidence_correlation_id": evidence_package_id,
		"playbook_version": evaluation["playbook_version"],
		"business_link": f"/coordination-cases/{case_id}",
	}
	with _CASE_REGISTER_LOCK:
		_CASE_REGISTER[case_id] = {
			"approved_payload": {
				"country": canonical_country,
				"evidence_package_id": evidence_package_id,
				"playbook_version": evaluation["playbook_version"],
				"reason": reason.strip(),
				"lead_agency": evaluation["lead_agency"],
				"participating_agencies": expected_agencies,
				"review_steps": expected_steps,
			},
			"receipt": receipt,
		}

	return dict(receipt)


LOCAL_FUNCTIONS: dict[str, Callable[..., dict[str, Any]]] = {
	"get_resilience_priorities": get_resilience_priorities,
	"get_country_resilience_evidence": get_country_resilience_evidence,
	"generate_resilience_report": generate_resilience_report,
	"evaluate_coordination_playbook": evaluate_coordination_playbook,
	"open_coordination_case": open_coordination_case,
}


def call_local_function(
	name: str,
	arguments: str | dict[str, Any],
	*,
	allow_side_effects: bool = False,
) -> str:
	"""Dispatch one Responses API function call and return JSON output."""
	try:
		payload = json.loads(arguments or "{}") if isinstance(arguments, str) else arguments
		if not isinstance(payload, dict):
			raise TypeError("function arguments must decode to a JSON object")

		function = LOCAL_FUNCTIONS.get(name)
		if function is None:
			raise ValueError(f"No local function is registered for {name!r}")
		if name == "open_coordination_case" and not allow_side_effects:
			raise PermissionError(
				"open_coordination_case requires explicit application approval"
			)

		return json.dumps(function(**payload), ensure_ascii=True)
	except (PermissionError, TypeError, ValueError, RuntimeError) as exc:
		return json.dumps(
			{
				"error": str(exc),
				"error_type": type(exc).__name__,
				"function": name,
			},
			ensure_ascii=True,
		)